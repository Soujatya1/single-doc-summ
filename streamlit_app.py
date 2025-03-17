import streamlit as st
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os
import pdfplumber
import pandas as pd
import extract_msg

BRD_FORMAT = """
## 1.0 Introduction
    ## 1.1 Purpose
    ## 1.2 To be process / High level solution
## 2.0 Impact Analysis
    ## 2.1 System impacts – Primary and cross functional
    ## 2.2 Impacted Products
    ## 2.3 List of APIs required
## 3.0 Process / Data Flow diagram / Figma
## 4.0 Business / System Requirement
    ## 4.1 Application / Module Name
    ## 4.2 Application / Module Name
## 5.0 MIS / DATA Requirement
## 6.0 Communication Requirement
## 7.0 Test Scenarios
## 8.0 Questions / Suggestions
## 9.0 Reference Document
## 10.0 Appendix
"""

@st.cache_resource
def initialize_llm():
    model = ChatGroq(
        groq_api_key="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", 
        model_name="Llama3-70b-8192"
    )
    
    llm_chain = LLMChain(
        llm=model, 
        prompt=PromptTemplate(
            input_variables=['requirements', 'tables', 'brd_format'],
            template="""
            Create a Business Requirements Document (BRD) based on the following details:

        Document Structure:
        {brd_format}

        Requirements:
        Analyze the content provided in the requirement documents and map the relevant information to each section defined in the BRD structure according to the {requirements}. Be concise and specific.

        Tables:
        If applicable, include the following tabular information extracted from the documents:
        {tables}

        Formatting:
        1. Use headings and subheadings for clear organization.
        2. Include bullet points or numbered lists where necessary for better readability.
        3. Clearly differentiate between functional and non-functional requirements.
        4. Provide tables in a well-structured format, ensuring alignment and readability.

        Key Points:
        1. Use the given format `{brd_format}` strictly as the base structure for the BRD.
        2. Ensure all relevant information from the requirements is displayed under the corresponding section.
        3. Avoid including irrelevant or speculative information.
        4. Summarize lengthy content while preserving its meaning.

        Output:
        The output must be formatted cleanly as a Business Requirements Document, following professional standards. Avoid verbose language and stick to the structure defined above.
        """
        )
    )
    return llm_chain

@st.cache_resource
def initialize_test_scenario_generator():
    model = ChatGroq(
        groq_api_key="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", 
        model_name="Llama3-70b-8192"
    )
    
    test_scenario_chain = LLMChain(
        llm=model, 
        prompt=PromptTemplate(
            input_variables=['brd_content'],
            template="""
            Based on the following Business Requirements Document (BRD), generate detailed test scenarios for section 7.0 Test Scenarios:
            
            BRD Content:
            {brd_content}
            
            Special Instructions for Test Scenarios Section:
            Based on the entire BRD content, generate at least 5 detailed test scenarios that would comprehensively validate the requirements. For each test scenario:
            - Provide a clear test ID and descriptive name
            - Include test objective/purpose
            - List detailed test steps
            - Define expected results/acceptance criteria
            - Specify test data requirements if applicable
            - Indicate whether it's a positive or negative test case
            - Note any dependencies or prerequisites
            """
        )
    )
    return test_scenario_chain

def extract_tables_from_excel(excel_file):
    """Extract tables from Excel file and return as formatted text"""
    tables_text = []
    
    try:
        # Read all sheets
        excel_data = pd.read_excel(excel_file, sheet_name=None)
        
        # Process each sheet
        for sheet_name, df in excel_data.items():
            if not df.empty:
                tables_text.append(f"Table from sheet '{sheet_name}':")
                tables_text.append(df.to_string(index=False))
                tables_text.append("\n")
    except Exception as e:
        st.error(f"Error processing Excel file: {str(e)}")
    
    return "\n".join(tables_text)

def extract_content_from_msg(msg_file):
    """Extract content from Outlook MSG file"""
    try:
        temp_file = BytesIO(msg_file.getvalue())
        temp_file.name = msg_file.name
        
        msg = extract_msg.Message(temp_file)
        
        content = []
        content.append(f"Subject: {msg.subject}")
        content.append(f"From: {msg.sender}")
        content.append(f"To: {msg.to}")
        content.append(f"Date: {msg.date}")
        content.append("\nBody:")
        content.append(msg.body)
        
        # Check for attachments
        if msg.attachments:
            content.append("\nAttachments mentioned (not processed):")
            for attachment in msg.attachments:
                content.append(f"- {attachment.longFilename}")
        
        return "\n".join(content)
    except Exception as e:
        st.error(f"Error processing MSG file: {str(e)}")
        return ""

st.title("Business Requirements Document Generator")

st.subheader("Document Logo")
logo_file = st.file_uploader("Upload logo/icon for document (PNG):", type=['png'])

if logo_file is not None:
    st.image(logo_file, caption="Logo Preview", width=100)
    st.success("Logo uploaded successfully! It will be added to the document header.")
    if 'logo_data' not in st.session_state:
        st.session_state.logo_data = logo_file.getvalue()
else:
    st.info("Please upload a PNG logo/icon that will appear in the document header.")

st.subheader("Requirement Documents")
uploaded_files = st.file_uploader("Upload requirement documents (PDF/DOCX/XLSX/MSG):", 
                                 accept_multiple_files=True, 
                                 type=['pdf', 'docx', 'xlsx', 'msg'])

if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = {'requirements': '', 'tables': ''}

if uploaded_files:
    combined_requirements = []
    all_tables_as_text = []
    
    for uploaded_file in uploaded_files:
        file_extension = os.path.splitext(uploaded_file.name)[-1].lower()
        st.write(f"Processing {uploaded_file.name}...")
        
        if file_extension == ".docx":
            doc = Document(uploaded_file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            combined_requirements.append(text)
            
            # Extract tables from DOCX
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                all_tables_as_text.append("\n".join(table_text))
        
        elif file_extension == ".pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
                combined_requirements.append(text)
                
                # Extract tables from PDF
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = []
                        for row in table:
                            row_text = [str(cell) if cell else "" for cell in row]
                            table_text.append(" | ".join(row_text))
                        all_tables_as_text.append("\n".join(table_text))
        
        elif file_extension == ".xlsx":
            excel_tables = extract_tables_from_excel(uploaded_file)
            all_tables_as_text.append(excel_tables)
            # Also add the tables as requirements text
            combined_requirements.append(f"Excel file content from {uploaded_file.name}:\n{excel_tables}")
        
        elif file_extension == ".msg":
            msg_content = extract_content_from_msg(uploaded_file)
            if msg_content:
                combined_requirements.append(msg_content)
        
        else:
            st.warning(f"Unsupported file format: {uploaded_file.name}")
    
    st.session_state.extracted_data = {
        'requirements': "\n\n".join(combined_requirements),
        'tables': "\n\n".join(all_tables_as_text)
    }

def add_header_with_logo(doc, logo_bytes):
    section = doc.sections[0]
    
    header = section.header
    
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    
    run = header_para.add_run()
    logo_stream = BytesIO(logo_bytes)
    run.add_picture(logo_stream, width=Inches(2.0))

if st.button("Generate BRD") and uploaded_files:
    if not st.session_state.extracted_data['requirements']:
        st.error("No content extracted from documents.")
    else:
        st.write("Generating BRD...")
        llm_chain = initialize_llm()
        
        prompt_input = {
            "requirements": st.session_state.extracted_data['requirements'],
            "tables": st.session_state.extracted_data['tables'],
            "brd_format": BRD_FORMAT
        }
        
        output = llm_chain.run(prompt_input)
        
        test_scenario_generator = initialize_test_scenario_generator()
        test_scenarios = test_scenario_generator.run({"brd_content": output})
        
        output = output.replace("7.0 Test Scenarios", "7.0 Test Scenarios\n" + test_scenarios)
        
        st.success("BRD generated successfully!")
        
        st.subheader("Generated Business Requirements Document")
        st.markdown(output)
        
        doc = Document()
        doc.add_heading('Business Requirements Document', level=0)
        
        if logo_file:
            logo_bytes = logo_file.getvalue()
            add_header_with_logo(doc, logo_bytes)
        
        doc.add_heading('Version History', level=1)
        version_table = doc.add_table(rows=1, cols=5)
        version_table.style = 'Table Grid'
        hdr_cells = version_table.rows[0].cells
        hdr_cells[0].text = 'Version'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Author'
        hdr_cells[3].text = 'Change description'
        hdr_cells[4].text = 'Review by'

        for _ in range(4):
            version_table.add_row()

        doc.add_paragraph('**Review by should be someone from IT function.**', style='Caption')

        doc.add_heading('Sign-off Matrix', level=1)
        signoff_table = doc.add_table(rows=1, cols=5)
        signoff_table.style = 'Table Grid'
        hdr_cells = signoff_table.rows[0].cells
        hdr_cells[0].text = 'Version'
        hdr_cells[1].text = 'Sign-off Authority'
        hdr_cells[2].text = 'Business Function'
        hdr_cells[3].text = 'Sign-off Date'
        hdr_cells[4].text = 'Email Confirmation'

        for _ in range(4):
            signoff_table.add_row()

        doc.add_page_break()

        doc.add_heading('Table of Contents', level=1)

        toc_paragraph = doc.add_paragraph()
        toc_paragraph.bold = True

        toc_entries = [
            "1.0 Introduction",
            "    1.1 Purpose",
            "    1.2 To be process / High level solution",
            "2.0 Impact Analysis",
            "    2.1 System impacts – Primary and cross functional",
            "    2.2 Impacted Products",
            "    2.3 List of APIs required",
            "3.0 Process / Data Flow diagram / Figma",
            "4.0 Business / System Requirement",
            "    4.1 Application / Module Name",
            "    4.2 Application / Module Name",
            "5.0 MIS / DATA Requirement",
            "6.0 Communication Requirement",
            "7.0 Test Scenarios",
            "8.0 Questions / Suggestions",
            "9.0 Reference Document",
            "10.0 Appendix"
        ]

        for entry in toc_entries:
            if entry.startswith("    "):
                doc.add_paragraph(entry.strip(), style='Heading 3')
            else:
                doc.add_paragraph(entry, style='Heading 2')

        doc.add_page_break()
        
        for section in output.split('\n#'):
            if not section.strip():
                continue
            
            lines = section.strip().split('\n')
            heading_text = lines[0].lstrip('#').strip()
            heading_level = 1 if section.startswith('#') else 2
            doc.add_heading(heading_text, level=heading_level)
            
            content = '\n'.join(lines[1:]).strip()
            if content:
                doc.add_paragraph(content)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Download BRD as Word document",
            data=buffer,
            file_name="Business_Requirements_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
